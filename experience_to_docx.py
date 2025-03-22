
from docx import Document
from lxml import html

class Experience2Docx:
    def __init__(self):
        self.doc = Document()

    def parse_html_string(self, html_content):
        # Parse the HTML content
        tree = html.fromstring(html_content)
        
        # Process each element in the parsed HTML
        for element in tree:
            if element.tag == 'ul':
                self.handle_unordered_list(element)
            elif element.tag == 'li':
                self.handle_list_item(element)
            else:
                self.handle_other_elements(element)
        
        return self.doc
    
    def handle_unordered_list(self, ul_element):
        # Create a bullet-point list in DOCX for <ul>
        for li in ul_element:
            if li.tag == 'li':
                self.handle_list_item(li)
    
    def handle_list_item(self, li_element):
        # Add a list item as a bullet point in DOCX for <li>
        paragraph = self.doc.add_paragraph(style='List Bullet')
        self.process_text(li_element, paragraph)
    
    def handle_other_elements(self, element):
        # Handle other elements like <strong> and <p>
        if element.tag == 'strong':
            # Bold text from <strong>
            paragraph = self.doc.add_paragraph()
            self.process_text(element, paragraph, bold=True)
        elif element.tag == 'p':
            # Regular paragraph from <p>
            paragraph = self.doc.add_paragraph()
            self.process_text(element, paragraph)

    def process_text(self, element, paragraph, bold=False):
        # Process the text inside an element (like <li>, <strong>, etc.)
        # If the element has nested tags (like <strong> inside <li>), handle those too
        for child in element.iter():
            if child.tag == 'strong':
                run = paragraph.add_run(child.text)
                run.bold = True
            else:
                run = paragraph.add_run(child.text)
                if bold:
                    run.bold = True