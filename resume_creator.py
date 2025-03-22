import io
import re
import os
import shutil
from html2docx import html2docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class ResumeCreator:
    def __init__(self, base_docx, destination_docx):
        self.base = base_docx
        self.dest = destination_docx
        
        os.makedirs("resume_part", exist_ok=True)
    
    def CnvHTML2ParaText(self, html_content, output_path):
        """
        Converts an HTML fragment into a formatted DOCX paragraph.
        """
        try:
            def extract_color_from_html(html):
                color_matches = re.findall(r'color\s*:\s*#([0-9a-fA-F]{6})', html)
                return color_matches

            full_html_content = f"<html><body>{html_content}</body></html>"
            buf = html2docx(full_html_content, title="Converted Document")

            doc = Document(io.BytesIO(buf.getvalue()))
            color_codes = extract_color_from_html(html_content)

            if color_codes:
                color_iter = iter(color_codes)
                for para in doc.paragraphs:
                    for run in para.runs:
                        color_code = next(color_iter, None)
                        if color_code:
                            rgb_color = RGBColor(int(color_code[:2], 16), int(color_code[2:4], 16), int(color_code[4:], 16))
                            run.font.color.rgb = rgb_color

            doc.save(output_path)
            return doc

        except Exception as e:
            print(f"Error during conversion or reading DOCX: {e}")
            return None

    def GenPlaceholderDocx(self, placeholder, style, source_path):
        """
        Replaces a placeholder in the destination DOCX file with content from a source DOCX file.
        Preserves formatting, fonts, colors, and sizes.
        """
        destination_doc = Document(self.dest)
        source_doc = Document(source_path)

        for i, para in enumerate(destination_doc.paragraphs):
            if placeholder in para.text:
                if 'List Bullet' not in [style.name for style in destination_doc.styles]:
                    bullet_style = destination_doc.styles.add_style('List Bullet', 1)  # 1 = Paragraph style
                    
                    # Customize the font of the style (you can modify as needed)
                    font = bullet_style.font
                    font.name = 'Calibri'
                    font.size = Pt(12)

                    # Access the underlying XML to adjust list behavior (bullets and indentation)
                    paragraph_format = bullet_style.paragraph_format
                    paragraph_format.left_indent = Pt(18)  # Indentation for bullets
                    paragraph_format.space_after = Pt(6)   # Space after each bullet

                    # Apply XML for bullet point styling
                    p = bullet_style._element  # Access the underlying XML element
                    numbering = OxmlElement('w:numPr')
                    num_id = OxmlElement('w:numId')
                    num_id.set(qn('w:val'), '1')  # You can customize the numId here
                    numbering.append(num_id)
                    p.insert(0, numbering)  # Insert the numbering XML into the paragraph
                
                for run in para.runs:
                    if placeholder in run.text:
                        for source_para in source_doc.paragraphs:
                            new_paragraph = destination_doc.add_paragraph(style=style)
                            new_paragraph.alignment = 0  # Align to left
                            for source_run in source_para.runs:
                                new_run = new_paragraph.add_run(source_run.text)
                                new_run.font.bold = source_run.font.bold
                                new_run.font.italic = source_run.font.italic
                                new_run.font.underline = source_run.font.underline
                                new_run.font.size = source_run.font.size
                                new_run.font.color.rgb = source_run.font.color.rgb
                    else:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                        new_run.font.underline = run.font.underline
                        new_run.font.size = run.font.size
                        new_run.font.color.rgb = run.font.color.rgb

                para._element.getparent().remove(para._element)
                # destination_doc.element.body.insert(i, new_paragraph._element)
                
                for new_para in destination_doc.paragraphs[-len(source_doc.paragraphs):]:
                    destination_doc.element.body.insert(i, new_para._element)
                break  

        destination_doc.save(self.dest)
        print(f"Replaced placeholders in '{self.dest}' with respective source documents.")

    def Gen(self, replacements):
        """
        Processes and replaces placeholders in a DOCX file with HTML-based formatted content.
        """
        shutil.copy2(self.base, self.dest)

        sources_dict = {}
        for placeholder, (style, html_content) in replacements.items():
            temp_docx_path = os.path.join("resume_part", f"{placeholder.strip('{}')}.docx")
            self.CnvHTML2ParaText(html_content, temp_docx_path)
            sources_dict[placeholder] = (style, temp_docx_path)

        for placeholder, (style, file_path) in sources_dict.items():
            self.GenPlaceholderDocx(placeholder, style, file_path)
